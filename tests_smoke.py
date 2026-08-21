"""Быстрая проверка ядра без сети: чтение файлов, чанкинг, векторная база.

Запуск: python tests_smoke.py
"""
import os
import shutil
import tempfile

from docx import Document
import pypdf

from config import CHUNK_OVERLAP, CHUNK_SIZE
from rag import chunking, loaders, store


class FakeEmbeddings:
    """Детерминированные векторы, чтобы не дёргать Ollama в тесте."""
    def __init__(self, dim=8):
        self.dim = dim

    def embed_documents(self, texts):
        return [[self._h(t, i) for i in range(self.dim)] for t in texts]

    def embed_query(self, text):
        return [self._h(text, i) for i in range(self.dim)]

    def _h(self, text, i):
        return sum(ord(c) for c in text) % 100 / 100.0 + i * 0.001


def make_pdf(path, pages):
    from reportlab.pdfgen import canvas
    c = canvas.Canvas(str(path))
    for i, p in enumerate(pages, 1):
        c.setPageSize((400, 400))
        c.drawString(72, 340, p)
        c.showPage()
    c.save()


def make_pdf_encrypted(path):
    # Создаём PDF и защищаем паролем через pypdf
    from reportlab.pdfgen import canvas
    tmp = path + ".tmp"
    c = canvas.Canvas(tmp)
    c.drawString(72, 340, "secret")
    c.showPage()
    c.save()
    reader = pypdf.PdfReader(tmp)
    writer = pypdf.PdfWriter()
    writer.append_pages_from_reader(reader)
    writer.encrypt(user_password="pass", owner_password="owner")
    with open(path, "wb") as fh:
        writer.write(fh)
    os.unlink(tmp)


def main():
    tmp = tempfile.mkdtemp()
    try:
        # --- TXT ---
        txt = os.path.join(tmp, "notes.txt")
        with open(txt, "w", encoding="cp1251") as fh:
            fh.write("Первая часть про отпуска.\n\n" + "Вторая часть про зарплату. " * 40)
        pages = loaders.read_document(txt)
        assert pages, "TXT не прочитан"
        print(f"TXT OK: {len(pages)} абзацев")

        # --- DOCX ---
        docx_path = os.path.join(tmp, "contract.docx")
        doc = Document()
        doc.add_paragraph("Пункт 1. Стороны договорились...")
        doc.add_paragraph("Пункт 2. Срок действия — один год.")
        doc.save(docx_path)
        pages = loaders.read_document(docx_path)
        assert len(pages) == 2
        print(f"DOCX OK: {len(pages)} абзацев")

        # --- PDF ---
        pdf = os.path.join(tmp, "report.pdf")
        make_pdf(pdf, ["Отчёт за квартал.", "Объём продаж вырос на 20%. " * 60])
        pages = loaders.read_document(pdf)
        assert len(pages) == 2, pages
        assert pages[0]["loc"] == "Стр. 1"
        print(f"PDF OK: {len(pages)} страниц")

        # --- Ошибки ---
        bad_pdf = os.path.join(tmp, "bad.pdf")
        with open(bad_pdf, "wb") as fh:
            fh.write(b"this is not a pdf at all")
        try:
            loaders.read_document(bad_pdf)
            raise SystemExit("Ошибка: битый PDF не отловлен")
        except loaders.LoadError as exc:
            print(f"Битый PDF OK: {exc}")

        enc_pdf = os.path.join(tmp, "enc.pdf")
        make_pdf_encrypted(enc_pdf)
        try:
            loaders.read_document(enc_pdf)
            raise SystemExit("Ошибка: зашифрованный PDF не отловлен")
        except loaders.LoadError as exc:
            print(f"Зашифрованный PDF OK: {exc}")

        empty_txt = os.path.join(tmp, "empty.txt")
        open(empty_txt, "w").close()
        try:
            loaders.read_document(empty_txt)
            raise SystemExit("Ошибка: пустой TXT не отловлен")
        except loaders.LoadError as exc:
            print(f"Пустой TXT OK: {exc}")

        try:
            loaders.read_document(os.path.join(tmp, "x.exe"))
            raise SystemExit("Ошибка: неподдерживаемый формат не отловлен")
        except loaders.LoadError as exc:
            print(f"Формат OK: {exc}")

        # --- Чанкинг ---
        chunks = chunking.split_pages(pages, CHUNK_SIZE, CHUNK_OVERLAP)
        assert chunks, "Чанкинг вернул пусто"
        assert all(c["source"] == "report.pdf" for c in chunks)
        print(f"Чанкинг OK: {len(chunks)} чанков")

        # --- Векторная база ---
        db = os.path.join(tmp, "db")
        vs = store.make_store(FakeEmbeddings(), db)
        n = store.add_chunks(vs, chunks)
        assert n == len(chunks)
        assert store.collection_size(vs) == n
        results = store.search(vs, "рост продаж", k=2)
        assert len(results) == 2
        print(f"База OK: добавили {n}, поиск вернул {len(results)}")
        print("known sources:", store.known_sources(vs))

        store.delete_source(vs, "report.pdf")
        assert store.collection_size(vs) == 0
        print("Удаление OK")

        print("\nВСЕ ПРОВЕРКИ ПРОШЛИ")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


if __name__ == "__main__":
    main()
